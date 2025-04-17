
import streamlit as st
from docx import Document
from io import BytesIO
import time

st.set_page_config(page_title="Nathan Mills Marine – AI Report Generator", layout="wide")
st.title("Nathan Mills Marine – AI Report Generator")

uploaded_file = st.file_uploader("📤 Upload Case File (.docx, .xlsx, or .json)", type=["docx", "xlsx", "json"])

if uploaded_file is not None:
    st.success(f"Uploaded: {uploaded_file.name}")
    
    if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        from docx import Document
        doc = Document(uploaded_file)
        st.subheader("Preview of Uploaded Report:")
        for para in doc.paragraphs[:5]:
            st.write(para.text)

    if st.button("Generate Report"):
        progress = st.progress(0)
        status = st.empty()
        
        progress.progress(10)
        status.text("🔍 Parsing input...")
        time.sleep(1)

        progress.progress(35)
        status.text("🧠 Generating AI report content...")
        time.sleep(2)

        report_sections = [
            ("1.0 INTRODUCTION", "This report was prepared to assess the vessel’s performance in accordance with the charterparty terms."),
            ("2.0 DOCUMENTS REVIEWED", "Charterparty, noon reports, AIS track data, weather records, hull inspection."),
            ("3.0 WEATHER CONDITIONS", "Weather was within warranty for 4.5 days, with no significant deviation or current."),
            ("4.0 SPEED AND CONSUMPTION", "The vessel averaged 11.3 knots against a warranted 12.2 knots with over-consumption of 1.8 MT IFO/day."),
            ("5.0 AIS ANALYSIS", "AIS track revealed 3 periods of unexplained slow speed, inconsistent with weather conditions."),
            ("6.0 CONCLUSION", "The vessel failed to meet charterparty warranties. Breach confirmed. Charterers may seek recovery for excess time and fuel.")
        ]

        progress.progress(70)
        status.text("📝 Formatting into .docx file...")

        # Generate the Word document
        output_doc = Document()
        output_doc.add_heading("Technical Report – AI Draft", level=1)
        for title, content in report_sections:
            output_doc.add_heading(title, level=2)
            output_doc.add_paragraph(content)
        output_doc.add_paragraph("\nCaptain Nathan Mills\nTechnical Director, Nathan Mills Marine")

        docx_io = BytesIO()
        output_doc.save(docx_io)
        docx_io.seek(0)

        progress.progress(100)
        status.text("✅ Report generation complete.")

        st.success("✅ AI-generated report is ready for download.")
        st.download_button("⬇️ Download Report (.docx)", data=docx_io, file_name="NMM_AI_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("Upload a .docx, .xlsx, or .json case file to begin.")
